import os, shutil, fitz
from docx import Document  #pip install pymupdf python-docx
from datetime import datetime
from pathlib import Path
from uuid import uuid4
from openai import OpenAI
import httpx
from fastapi import Depends, Query, HTTPException, Form, UploadFile, File, Body, APIRouter, Header, Request
from utils.helpers import is_user_in_workspace, generate_api_key
from sqlalchemy.orm import Session
import uuid 
from fastapi.responses import JSONResponse
import requests
from bs4 import BeautifulSoup
import time,json
from livekit import api
from typing import List, Optional
from twilio.rest import Client

from sqlalchemy.orm import sessionmaker, Session
from db.database import get_db, get_current_user
from models.models import ( User, Workspace, WorkspaceSettings, WorkspaceMember, KnowledgeBase, KnowledgeFile, APIKey, FileStatus, SourceStatus, 
                           pbx_ai_agent, PBXLLM, ChatSession, LLMVoice, ImportedPhoneNumber, WebCall, PhoneNumber
                           )
from models.schemas import ( UserSignup, UserLogin, UpdateUser,DispatchRequest,CreateRoomRequestSchema,WebCallResponse,WebCallCreateRequest,GetPBXLLMOut,
                                WorkspaceCreate, WorkspaceOut, InviteMember,WorkspaceSettingsUpdate, AgentCreate, AgentOut, PBXLLMCreate, PBXLLMOut, 
                                CreateChatRequest, CreateChatResponse, VoiceOut, VoiceCreate, PhoneNumberCreate, PhoneNumberOut, APIResponse, PhoneNumberRequest, 
                                Response, AgentUpdate
                     )
from utils.helpers import (
    format_response, validate_email,
    validate_password, format_datetime_ist
)
from utils.security import (
    hash_password, verify_password,
    create_token
)

from starlette.config import Config

from models.twilio import PhoneNumberOut

from db.database import engine

router = APIRouter()

SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "uploads")).resolve()
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

@router.get("/check-username")
def check_username_availability(
    username: str = Query(..., min_length=5, max_length=50),
    db: Session = Depends(get_db)
):
    try:
        existing_user = db.query(User).filter(User.username == username).first()
        if existing_user:
            return format_response(
                status=False,
                message="Username already taken",
                errors=[{"field": "username", "message": "This username is already in use"}]
            )
        return format_response(
            status=True,
            message="Username is available"
        )

    except Exception as e:
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.post("/signup") 
def signup(user: UserSignup, db: Session = Depends(get_db)):
    try:
        # Validate inputs
        validate_email(user.email)
        validate_password(user.password)

        # Check if user already exists
        if db.query(User).filter(User.username == user.username).first():
            return format_response(
                status=False,
                message="Username already exists",
                errors=[{"field": "username", "message": "Username already exists"}]
            )

        if db.query(User).filter(User.email == user.email).first():
            return format_response(
                status=False,
                message="Email already exists",
                errors=[{"field": "email", "message": "Email already exists"}]
            )

        # Create user
        new_user = User(
            username=user.username,
            email=user.email,
            full_name=user.full_name,
            hashed_password=hash_password(user.password),
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        db.add(new_user)
        db.commit()
        db.refresh(new_user)

        # Create default workspace
        workspace = Workspace(
            name=f"{user.username}_workspace",
            description="Default workspace",
            owner_id=new_user.id,
            created_at=datetime.utcnow()
        )
        db.add(workspace)
        db.commit()
        db.refresh(workspace)

        # Add user as member and create workspace settings
        settings = WorkspaceSettings(
            workspace_id=workspace.id,
            default_model="gpt-4",
            default_voice="echo",
            temperature=1
        )
        member = WorkspaceMember(
            user_id=new_user.id,
            workspace_id=workspace.id,
            role="owner"
        )
        db.add_all([settings, member])
        db.commit()

        return format_response(
            status=True,
            message="Signup successful",
            data={
                "workspace_id": workspace.id,
                "workspace_name": workspace.name
            }
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.post("/login")
def login(user: UserLogin, db: Session = Depends(get_db)):
    try:
        db_user = db.query(User).filter(User.email == user.email).first()
        if not db_user or not verify_password(user.password, db_user.hashed_password):
            raise HTTPException(status_code=401, detail="Invalid credentials")

        token = create_token(data={"email": db_user.email})

        return format_response(
            status=True,
            message="Login successful",
            data={
                "token": 'Bearer ' + token
            }
        )

    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")

@router.patch("/user/update")
def update_user(
    user_data: UpdateUser,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        existing_user = db.query(User).filter(User.id == current_user.id).first()
        if not existing_user:
            return format_response(
                status=False,
                message="User not found",
                errors=[{"field": "user", "message": "User not found"}]
            )

        updated = False
        errors = []

        # Update username
        if user_data.username and user_data.username != existing_user.username:
            username_exists = db.query(User).filter(User.username == user_data.username).first()
            if username_exists:
                errors.append({
                    "field": "username",
                    "message": "Username already taken"
                })
            else:
                existing_user.username = user_data.username
                updated = True

        # Update email
        if user_data.email and user_data.email != existing_user.email:
            try:
                validate_email(user_data.email)
            except Exception as e:
                errors.append({
                    "field": "email",
                    "message": str(e)
                })
            else:
                email_exists = db.query(User).filter(User.email == user_data.email).first()
                if email_exists:
                    errors.append({
                        "field": "email",
                        "message": "Email already registered"
                    })
                else:
                    existing_user.email = user_data.email
                    updated = True

        # Update password
        if user_data.password:
            try:
                validate_password(user_data.password)
            except Exception as e:
                errors.append({
                    "field": "password",
                    "message": str(e)
                })
            else:
                existing_user.hashed_password = hash_password(user_data.password)
                updated = True

        if errors:
            return format_response(
                status=False,
                message="Validation or conflict error",
                errors=errors
            )

        if not updated:
            return format_response(
                status=False,
                message="No valid fields provided for update",
                errors=[{"field": "update", "message": "Nothing to update"}]
            )

        # Update the timestamp
        existing_user.updated_at = datetime.utcnow()

        db.commit()
        user_response = {
            "id": existing_user.id,
            "username": existing_user.username,
            "email": existing_user.email,
            "full_name": existing_user.full_name,
            # Add other fields you want to return here
        }

        return format_response(
            status=True,
            message="User details updated successfully",
            data=user_response
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.post("/create-room-token")
async def create_room_and_token(
    request: CreateRoomRequestSchema,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)  # get the logged-in user
):
    try:
        LIVEKIT_API_KEY = os.getenv("LIVEKIT_API_KEY")
        LIVEKIT_API_SECRET = os.getenv("LIVEKIT_API_SECRET")
        LIVEKIT_URL = os.getenv("LIVEKIT_URL")

        if not (LIVEKIT_API_KEY and LIVEKIT_API_SECRET and LIVEKIT_URL):
            raise HTTPException(status_code=500, detail="LiveKit credentials not configured")

        workspace_ids = [
            membership.workspace_id for membership in current_user.memberships
        ]

        if not workspace_ids:
            raise HTTPException(status_code=403, detail="User is not a member of any workspace")

        # 🔥 Step 2: Search for agent in user's workspaces
        agent = (
            db.query(pbx_ai_agent)
            .filter(
                pbx_ai_agent.name == request.agent_name,
                pbx_ai_agent.workspace_id.in_(workspace_ids)
            )
            .first()
        )
        pbxllm = (
            db.query(PBXLLM)
            .filter(PBXLLM.workspace_id.in_(workspace_ids))
            .first()
        )
        

        if not agent and pbxllm:
            raise HTTPException(status_code=404, detail="No LLM configuration & Agent not found for this user")
        

        # 🔥 Step 2: Create room (optional - LiveKit auto-creates)
        lkapi = api.LiveKitAPI(
            url=LIVEKIT_URL,
            api_key=LIVEKIT_API_KEY,
            api_secret=LIVEKIT_API_SECRET
        )
        try:
            lkapi.room.create_room(
                api.CreateRoomRequest(
                    name=request.room_name,
                    empty_timeout=10 * 60,  # 10 minutes timeout
                    max_participants=10,
                    metadata=json.dumps(request.metadata) if request.metadata else None
                )
            )
            print(f"Room '{request.room_name}' created successfully.")
        except Exception as e:
            if "already exists" not in str(e):
                raise HTTPException(status_code=500, detail=f"LiveKit Error: {e}")

        # 🔥 Step 3: Generate token
        token = api.AccessToken(LIVEKIT_API_KEY, LIVEKIT_API_SECRET) \
            .with_identity(request.participant_identity) \
            .with_name(request.participant_name) \
            .with_grants(api.VideoGrants(
                room_join=True,
                room=request.room_name,
            ))
        jwt_token = token.to_jwt()
        return jwt_token

    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")

@router.post("/create-dispatch")
async def create_dispatch(request: DispatchRequest):
    lkapi = api.LiveKitAPI()
    try:
        dispatch = await lkapi.agent_dispatch.create_dispatch(
            api.CreateAgentDispatchRequest(
                agent_name=request.agent_name,
                room=request.room_name,
                metadata=json.dumps(request.metadata)
            )
        )
        return {
            "status": "success",
            "dispatch_id": dispatch.id,
            "room": dispatch.room,
            "agent_name": dispatch.agent_name
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        await lkapi.aclose()    

@router.get("/user/status")
def check_user_status(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        user = db.query(User).filter(User.id == current_user.id).first()

        if not user:
            return format_response(
                status=False,
                message="User not found",
                errors=[{"field": "user", "message": "User not found"}]
            )

        return format_response(
            status=True,
            message="User is active" if user.is_active else "User is inactive",
            data={"is_active": user.is_active}
        )

    except Exception as e:
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.post("/workspaces")
def create_workspace(data: WorkspaceCreate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        workspace = Workspace(
            name=data.name,
            description=data.description,
            owner_id=current_user.id
        )
        db.add(workspace)
        db.commit()
        db.refresh(workspace)

        settings = WorkspaceSettings(
            workspace_id=workspace.id,
            default_model="gpt-4",
            default_voice="echo",
            temperature=1
        )
        db.add(settings)

        member = WorkspaceMember(
            user_id=current_user.id,
            workspace_id=workspace.id,
            role="owner"
        )
        db.add(member)
        db.commit()

        return format_response(
            status=True,
            message="Workspace created",
            data=WorkspaceOut(
                id=workspace.id,
                name=workspace.name,
                description=workspace.description,
                created_at=format_datetime_ist(workspace.created_at),
                owner_username=current_user.username,
            ).dict()
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.get("/workspaces")
def list_workspaces(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        workspaces = (
            db.query(Workspace)
            .join(WorkspaceMember)
            .filter(WorkspaceMember.user_id == current_user.id)
            .all()
        )

        data = [
            WorkspaceOut(
                id=w.id,
                name=w.name,
                created_at=format_datetime_ist(w.created_at),
                owner_username=w.owner.username
            ).dict() for w in workspaces
        ]

        return format_response(
            status=True,
            message="Workspaces retrieved successfully",
            data=data
        )

    except Exception as e:
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.post("/workspaces/{workspace_id}/invite")
def invite_user(
    workspace_id: int,
    invite: InviteMember,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
        if not workspace:
            return format_response(
                status=False,
                message="Workspace not found",
                errors=[{"field": "workspace_id", "message": "Workspace not found"}]
            )

        if workspace.owner_id != current_user.id:
            return format_response(
                status=False,
                message="Only the owner can invite members",
                errors=[{"field": "permission", "message": "Only the owner can invite"}],
                status_code=403
            )

        user = db.query(User).filter(User.username == invite.username).first()
        if not user:
            return format_response(
                status=False,
                message="User not found",
                errors=[{"field": "username", "message": "User not found"}]
            )

        existing = db.query(WorkspaceMember).filter_by(workspace_id=workspace_id, user_id=user.id).first()
        if existing:
            return format_response(
                status=False,
                message="User already a member",
                errors=[{"field": "username", "message": "Already a member"}],
                status_code=400
            )

        new_member = WorkspaceMember(workspace_id=workspace_id, user_id=user.id, role=invite.role)
        db.add(new_member)
        db.commit()

        return format_response(
            status=True,
            message=f"{invite.username} invited to workspace"
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.delete("/workspaces/{workspace_id}")
def delete_workspace(
    workspace_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
        if not workspace:
            return format_response(
                status=False,
                message="Workspace not found",
                errors=[{"field": "workspace_id", "message": "Workspace not found"}],
                status_code=404
            )

        if workspace.owner_id != current_user.id:
            return format_response(
                status=False,
                message="Only the owner can delete the workspace",
                errors=[{"field": "permission", "message": "Unauthorized"}],
                status_code=403
            )

        # Delete knowledge base folders from filesystem
        for kb in workspace.knowledge_bases:
            kb_path = UPLOAD_DIR / str(workspace.id) / str(kb.id)
            if kb_path.exists():
                import shutil
                shutil.rmtree(kb_path)

        # Delete workspace (with cascading)
        db.delete(workspace)
        db.commit()

        return format_response(
            status=True,
            message="Workspace deleted successfully"
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.post("/workspaces/knowledge-bases")
def create_knowledge_base(
    workspace_id: int = Form(...),
    name: str = Form(...),
    files: List[UploadFile] = File(default=[]),
    urls: Optional[List[str]] = Form(default=[]),
    text_filenames: Optional[List[str]] = Form(default=[]),
    text_contents: Optional[List[str]] = Form(default=[]),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # Validate workspace
        workspace = db.query(Workspace).filter(
            Workspace.id == workspace_id,
            Workspace.owner_id == current_user.id
        ).first()
        if not workspace:
            return format_response(
                status=False,
                message="Workspace not found",
                errors=[{"field": "workspace_id", "message": "Workspace not found"}],
                status_code=404
            )

        # Validate name
        if not name.strip():
            return format_response(
                status=False,
                message="Validation error",
                errors=[{"field": "name", "message": "Knowledge base name is required"}],
                status_code=400
            )

        # Create knowledge base folder and metadata
        kb_id = f"knowledge_base_{uuid.uuid4().hex[:16]}"
        kb_folder = UPLOAD_DIR / f"workspace_{workspace_id}" / kb_id
        kb_folder.mkdir(parents=True, exist_ok=True)
        now_utc = datetime.utcnow()

        # Create KnowledgeBase entry
        kb = KnowledgeBase(
            id=kb_id,
            name=name,
            file_path=str(kb_folder),
            workspace_id=workspace_id,
            enable_auto_refresh=True,
            auto_refresh_interval=24,
            last_refreshed=now_utc
        )
        db.add(kb)
        db.flush()  # Get kb.id before commit

        sources = []

        # Process uploaded files
        for file in files:
            if file and file.filename:
                filename = f"{uuid.uuid4().hex}_{file.filename}"
                file_path = kb_folder / filename
                with open(file_path, "wb") as f:
                    f.write(file.file.read())

                db.add(KnowledgeFile(
                    kb_id=kb.id,
                    filename=filename,
                    file_path=str(file_path),
                    source_type=SourceStatus.file,
                    status=FileStatus.pending
                ))

                sources.append({
                    "type": "document",
                    "source_id": kb.id,
                    "filename": filename,
                    "file_url": str(file_path)
                })

        # Process web URLs
        for url in urls or []:
            if url.strip():
                filename = f"web_{uuid.uuid4().hex}.txt"
                file_path = kb_folder / filename
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(f"Crawled content from {url}")  # placeholder

                db.add(KnowledgeFile(
                    kb_id=kb.id,
                    filename=filename,
                    file_path=str(file_path),
                    source_type=SourceStatus.web_page,
                    status=FileStatus.pending
                ))

                sources.append({
                    "type": "web_page",
                    "source_id": kb.id,
                    "filename": filename,
                    "file_url": str(file_path)
                })

        # Process text entries
        for title, content in zip(text_filenames or [], text_contents or []):
            if title.strip() and content.strip():
                filename = f"{uuid.uuid4().hex}_{title}.txt"
                file_path = kb_folder / filename
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)

                db.add(KnowledgeFile(
                    kb_id=kb.id,
                    filename=filename,
                    file_path=str(file_path),
                    source_type=SourceStatus.text,
                    status=FileStatus.pending
                ))

                sources.append({
                    "type": "text",
                    "source_id": kb.id,
                    "filename": filename,
                    "file_url": str(file_path)
                })

        if not sources:
            return format_response(
                status=False,
                message="No valid sources provided",
                errors=[{"field": "sources", "message": "Provide at least one file, URL, or text"}]
            )

        db.commit()
        db.refresh(kb)

        return format_response(
            status=True,
            message="Knowledge base created successfully",
            data={
                "knowledge_base_id": kb.id,
                "knowledge_base_name": kb.name,
                "status": "in_progress",
                "knowledge_base_sources": sources,
                "enable_auto_refresh": kb.enable_auto_refresh,
                "last_refreshed_timestamp": int(now_utc.timestamp() * 1000)
            }
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )
    
@router.post("/workspaces/{workspace_id}/create-knowledge-base")
async def create_or_update_knowledge_base(
    workspace_id: int,
    kb_id: Optional[str] = Form(None),
    knowledge_base_name: Optional[str] = Form(None),
    knowledge_base_files: List[UploadFile] = File(default=[]),
    knowledge_base_urls: Optional[str] = Form(default=None),
    knowledge_base_texts: Optional[str] = Form(default=None),

    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        if knowledge_base_urls:
            try:
                knowledge_base_urls = json.loads(knowledge_base_urls)
            except json.JSONDecodeError:
                return format_response(
                    status=False,
                    message="Invalid JSON for knowledge_base_urls",
                    errors=[{"field": "knowledge_base_urls", "message": "Invalid JSON format"}],
                    status_code=400
                )
        else:
            knowledge_base_urls = []
        if knowledge_base_texts:
            try:
                knowledge_base_texts = json.loads(knowledge_base_texts)
            except json.JSONDecodeError:
                return format_response(
                    status=False,
                    message="Invalid JSON for knowledge_base_texts",
                    errors=[{"field": "knowledge_base_texts", "message": "Invalid JSON format"}],
                    status_code=400
                )
        else:
            knowledge_base_texts = []

        # Check workspace ownership
        workspace = db.query(Workspace).filter(
            Workspace.id == workspace_id
        ).first()
        if not workspace:
            return format_response(
                status=False,
                message="Workspace not found",
                errors=[{"field": "workspace_id", "message": "Workspace not found"}],
                status_code=404
            )

        # Check if KB exists
        if kb_id:
            # Try to fetch existing KB by kb_id
            kb = db.query(KnowledgeBase).filter_by(
                workspace_id=workspace_id, id=kb_id
            ).first()
        else:
            kb = None
       
        now_utc = datetime.utcnow()

        if not kb:
            # Create new KB
            kb_id = f"{uuid.uuid4().hex[:16]}"
            kb_folder = UPLOAD_DIR / f"workspace_{workspace_id}" / kb_id
            kb_folder.mkdir(parents=True, exist_ok=True)

            kb = KnowledgeBase(
                id=kb_id,
                name=knowledge_base_name,
                file_path=str(kb_folder),
                workspace_id=workspace_id,
                enable_auto_refresh=True,
                auto_refresh_interval=24,
                last_refreshed=now_utc
            )
            db.add(kb)
            db.flush()
        else:
            # Existing KB folder
            kb_folder = Path(kb.file_path)
            kb_folder.mkdir(parents=True, exist_ok=True)

        sources = []

        # Process uploaded files
        for file in knowledge_base_files:
            if file and file.filename:
                filename = f"{uuid.uuid4().hex}_{file.filename}"
                file_path = kb_folder / filename
                content = await file.read()
                with open(file_path, "wb") as f_out:
                    f_out.write(content)

                # Extract text for embedding
                ext = Path(filename).suffix.lower()
                extract_text = ""
                if ext == ".pdf":
                    with fitz.open(file_path) as doc:
                        extract_text = "\n".join(p.get_text() for p in doc)
                elif ext == ".docx":
                    extract_text = "\n".join(p.text for p in Document(file_path).paragraphs)
                elif ext == ".txt":
                    extract_text = content.decode("utf-8")
                else:
                    extract_text = ""  # fallback for unsupported

                # Truncate and embed
                truncated = extract_text[:30000]
                embedding_response = openai_client.embeddings.create(
                    model="text-embedding-3-small",
                    input=truncated
                )
                embedding_vector = embedding_response.data[0].embedding

                kb_file = KnowledgeFile(
                    kb_id=kb.id,
                    filename=filename,
                    file_path=str(file_path),
                    extract_data=extract_text,
                    embedding=embedding_vector,
                    status=FileStatus.completed,
                    source_type=SourceStatus.file  # ✅ updated to new enum
                )
                db.add(kb_file)
                db.flush() 
                sources.append({
                    "type": "document",
                    "source_id": kb_file.id,
                    "filename": filename,
                    "file_url": str(file_path)
                })

        # Process web URLs
        for url in knowledge_base_urls or []:
            if url.strip():
                filename = f"web_{uuid.uuid4().hex}.txt"
                file_path = kb_folder / filename
                res = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
                res.raise_for_status()
                soup = BeautifulSoup(res.text, "html.parser")
                for tag in soup(["script", "style"]): tag.decompose()
                extract_text = soup.get_text(separator="\n", strip=True)
                file_path.write_text(extract_text, encoding="utf-8")

                truncated = extract_text[:30000]
                embedding_response = openai_client.embeddings.create(
                    model="text-embedding-3-small",
                    input=truncated
                )
                embedding_vector = embedding_response.data[0].embedding

                kb_file = KnowledgeFile(
                    kb_id=kb.id,
                    filename=filename,
                    file_path=str(file_path),
                    extract_data=extract_text,
                    embedding=embedding_vector,
                    status=FileStatus.completed,
                    source_type=SourceStatus.url  # ✅ updated to new enum
                )
                db.add(kb_file)
                db.flush() 
                sources.append({
                    "type": "url",
                    "source_id": kb_file.id,
                    "filename": filename,
                    "file_url": str(file_path)
                })

        for entry in knowledge_base_texts or []:
            title = entry.get("title", "").strip()
            content = entry.get("text", "").strip()
            if title and content:
                filename = f"{uuid.uuid4().hex}_{title}.txt"
                file_path = kb_folder / filename
                file_path.write_text(content, encoding="utf-8")

                truncated = content[:30000]
                embedding_response = openai_client.embeddings.create(
                    model="text-embedding-3-small",
                    input=truncated
                )
                embedding_vector = embedding_response.data[0].embedding

                kb_file = KnowledgeFile(
                    kb_id=kb.id,
                    filename=filename,
                    file_path=str(file_path),
                    extract_data=content,
                    embedding=embedding_vector,
                    status=FileStatus.completed,
                    source_type=SourceStatus.txt  # ✅ updated to new enum
                )
                db.add(kb_file)
                db.flush() 
                sources.append({
                    "type": "text",
                    "source_id": kb_file.id,
                    "filename": filename,
                    "file_url": str(file_path)
                })

        if not sources:
            return format_response(
                status=False,
                message="No valid sources provided",
                errors=[{"field": "sources", "message": "Provide at least one file, URL, or text"}]
            )

        kb.last_refreshed = now_utc
        db.commit()
        db.refresh(kb)

        return format_response(
            status=True,
            message="Knowledge base updated successfully",
            data={
                "knowledge_base_id": kb.id,
                "knowledge_base_name": kb.name,
                "status": "in_progress",
                "knowledge_base_sources": sources,
                "enable_auto_refresh": kb.enable_auto_refresh,
                "last_refreshed_timestamp": int(now_utc.timestamp() * 1000)
            }
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}],
            status_code=500
        )


@router.get("/workspaces/{workspace_id}/list-knowledge-bases")
def list_knowledge_bases(
    workspace_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # Check workspace ownership or membership
        workspace = db.query(Workspace).filter(
            Workspace.id == workspace_id
        ).first()

        if not workspace:
            return format_response(
                status=False,
                message="Workspace not found",
                errors=[{"field": "workspace_id", "message": "Workspace not found"}],
                status_code=404
            )

        # Fetch all KnowledgeBases for workspace
        kbs = db.query(KnowledgeBase).filter_by(workspace_id=workspace_id).all()
       
        response_data = []
        for kb in kbs:
            # Fetch all files (KnowledgeFile) for this KB
            knowledge_files = db.query(KnowledgeFile).filter_by(kb_id=kb.id).all()

            sources = []
            for kf in knowledge_files:
                mapped_type = (
                    "document" if kf.source_type == "file"
                    else "text" if kf.source_type == "txt"
                    else kf.source_type
                )
                sources.append({
                    "type": mapped_type,  # file, url, txt
                    "source_id": kf.id,
                    "filename": kf.filename,
                    "file_url": kf.file_path
                })
               

            response_data.append({
                "knowledge_base_id": kb.id,
                "knowledge_base_name": kb.name,
                "status": "completed",  # You can make dynamic if needed
                "knowledge_base_sources": sources,
                "enable_auto_refresh": kb.enable_auto_refresh,
                "last_refreshed_timestamp": int(kb.last_refreshed.timestamp() * 1000) if kb.last_refreshed else None
            })

        return format_response(
            status=True,
            message="Knowledge bases fetched successfully",
            data=response_data
        )

    except Exception as e:
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}],
            status_code=500
        )

    
@router.get("/get-knowledge-base/{kb_id}")
def get_knowledge_base(
    kb_id: str,  # changed from int to str
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
        if not kb:
            return format_response(
                status=False,
                message="Knowledge base not found",
                errors=[{"field": "kb_id", "message": "Knowledge base not found"}],
                status_code=404
            )

        is_member = db.query(WorkspaceMember).filter_by(
            workspace_id=kb.workspace_id,
            user_id=current_user.id
        ).first()

        if not is_member:
            return format_response(
                status=False,
                message="Access denied",
                errors=[{"field": "workspace", "message": "Access denied"}],
                status_code=403
            )

        files = db.query(KnowledgeFile).filter_by(kb_id=kb.id).all()
        
        sources = [
            {
                "type": "document" if f.source_type == "file"
                        else "text" if f.source_type == "txt"
                        else f.source_type,  # map source_type to string
                "source_id": f.id,
                "filename": f.filename,
                "file_url": f.file_path
            } for f in files
        ]

        data = {
            "knowledge_base_id": kb.id,
            "knowledge_base_name": kb.name,
            "status": "in_progress",
            "knowledge_base_sources": sources,
            "enable_auto_refresh": kb.enable_auto_refresh,
            "last_refreshed_timestamp": int(kb.last_refreshed.timestamp() * 1000) if kb.last_refreshed else None
        }

        return format_response(
            status=True,
            message="Knowledge base details fetched",
            data=data
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}],
            status_code=500
        )
       
from openai import OpenAI

config = Config(".env")

# Ensure you initialize OpenAI client
openai_client = OpenAI()

@router.post("/knowledge-bases/sources")
async def add_kb_source(
    kb_id: str = Form(...),
    source_type: str = Form(...),
    file: UploadFile = File(None),
    url: str = Form(None),
    text_filename: str = Form(None),
    text_content: str = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # Validate KB and user membership
        kb = db.query(KnowledgeBase).filter_by(id=kb_id).first()
        if not kb:
            return format_response(False, "Knowledge base not found", errors=[{"field": "kb_id"}], status_code=404)

        member = db.query(WorkspaceMember).filter_by(workspace_id=kb.workspace_id, user_id=current_user.id).first()
        if not member:
            return format_response(False, "Access denied", errors=[{"field": "workspace"}], status_code=403)

        # Directory setup
        kb_path = UPLOAD_DIR / f"workspace_{kb.workspace_id}" / f"knowledge_base_{kb_id}"
        kb_path.mkdir(parents=True, exist_ok=True)

        # Process content
        file_path, file_name, extract_text = None, None, ""

        if source_type == "file":
            if not file:
                return format_response(False, "File is required", errors=[{"field": "file"}], status_code=400)

            ext = Path(file.filename).suffix.lower()
            file_name = f"{uuid4().hex}_{file.filename}"
            file_path = kb_path / file_name

            with open(file_path, "wb") as f_out:
                f_out.write(await file.read())

            if ext == ".pdf":
                with fitz.open(file_path) as doc:
                    extract_text = "\n".join(p.get_text() for p in doc)
            elif ext == ".docx":
                extract_text = "\n".join(p.text for p in Document(file_path).paragraphs)
            elif ext == ".txt":
                extract_text = file_path.read_text(encoding="utf-8")
            else:
                return format_response(False, "Unsupported file type", status_code=400)

        elif source_type == "web_page":
            if not url or not text_filename:
                return format_response(False, "URL and filename required", status_code=400)
            res = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
            res.raise_for_status()
            soup = BeautifulSoup(res.text, "html.parser")
            for tag in soup(["script", "style"]): tag.decompose()
            extract_text = soup.get_text(separator="\n", strip=True)
            file_name = text_filename
            file_path = url

        elif source_type == "txt":
            if not text_filename or not text_content:
                return format_response(False, "Text filename and content required", status_code=400)
            file_name = f"{uuid4().hex}_{text_filename}.txt"
            file_path = kb_path / file_name
            file_path.write_text(text_content, encoding="utf-8")
            extract_text = text_content

        else:
            return format_response(False, "Invalid source type", status_code=400)

        clean_text = extract_text.replace("\n", " ").strip()
        truncated = clean_text[:3000]  # limit tokens for embedding

        embedding_response = openai_client.embeddings.create(
            model="text-embedding-3-small",
            input=truncated
        )
        embedding_vector = embedding_response.data[0].embedding

        # Save to DB
        kb_file = KnowledgeFile(
            kb_id=kb.id,
            filename=file_name,
            file_path=str(file_path),
            extract_data=extract_text,
            embedding=embedding_vector,  # ➕ store embedding
            status=FileStatus.completed,
            source_type=SourceStatus(source_type)
        )
        db.add(kb_file)
        db.commit()

        return format_response(True, "Source added with embedding", data={
            "source_id": kb_file.id,
            "filename": file_name,
            "file_url": str(file_path),
            "type": source_type,
        })

    except Exception as e:
        db.rollback()
        return format_response(False, "Internal Server Error", errors=[{"field": "server", "message": str(e)}], status_code=500)

@router.delete("/knowledge-bases/{kb_id}")
def delete_knowledge_base(
    kb_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # Fetch knowledge base
        kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
        if not kb:
            return format_response(
                status=False,
                message="Knowledge base not found",
                errors=[{"field": "kb_id", "message": "Knowledge base not found"}],
                status_code=404
            )

        # Check workspace ownership
        workspace = db.query(Workspace).filter(Workspace.id == kb.workspace_id).first()
        if not workspace or workspace.owner_id != current_user.id:
            return format_response(
                status=False,
                message="Only the workspace owner can delete this knowledge base",
                errors=[{"field": "permission", "message": "Access denied"}],
                status_code=403
            )

        # Delete related knowledge files
        db.query(KnowledgeFile).filter(KnowledgeFile.kb_id == kb_id).delete()

        # Delete knowledge base
        db.delete(kb)
        db.commit()

        # Remove uploaded files from disk
        kb_path = UPLOAD_DIR / f"workspace_{workspace.id}" / f"knowledge_base_{kb.id}"
        if kb_path.exists():
            shutil.rmtree(kb_path)

        return format_response(
            status=True,
            message="Knowledge base deleted successfully"
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}],
            status_code=500
        )
    
@router.delete("/knowledge-files/{file_id}")
def delete_knowledge_file(
    file_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # Fetch the file entry
        file = db.query(KnowledgeFile).filter(KnowledgeFile.id == file_id).first()
        if not file:
            return format_response(
                status=False,
                message="File not found",
                errors=[{"field": "file_id", "message": "Knowledge file not found"}],
                status_code=404
            )

        # Check if user has access to the knowledge base
        kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == file.kb_id).first()
        if not kb:
            return format_response(
                status=False,
                message="Knowledge base not found",
                errors=[{"field": "kb_id", "message": "Associated knowledge base not found"}],
                status_code=404
            )

        membership = db.query(WorkspaceMember).filter_by(
            workspace_id=kb.workspace_id,
            user_id=current_user.id
        ).first()

        if not membership:
            return format_response(
                status=False,
                message="Access denied",
                errors=[{"field": "workspace", "message": "User is not a member of this workspace"}],
                status_code=403
            )

        # Attempt to delete the file from the filesystem
        file_path = Path(file.file_path)
        if file_path.exists() and file_path.is_file():
            try:
                file_path.unlink()
            except Exception as e:
                # Log warning or handle non-fatal file deletion error if needed
                pass

        # Delete the file record from the database
        db.delete(file)
        db.commit()

        return format_response(
            status=True,
            message="Knowledge file deleted successfully"
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}],
            status_code=500
        )
    
@router.get("/workspaces/{workspace_id}/settings")
def get_workspace_settings(
    workspace_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        membership = db.query(WorkspaceMember).filter_by(workspace_id=workspace_id, user_id=current_user.id).first()
        if not membership:
            return format_response(
                status=False,
                message="Access denied",
                errors=[{"field": "workspace_id", "message": "User is not a member of this workspace"}],
                status_code=403
            )

        settings = db.query(WorkspaceSettings).filter_by(workspace_id=workspace_id).first()
        if not settings:
            return format_response(
                status=False,
                message="Settings not found",
                errors=[{"field": "workspace_id", "message": "Workspace settings not found"}],
                status_code=404
            )

        return format_response(
            status=True,
            message="Workspace settings retrieved",
            data={
                "default_voice": settings.default_voice,
                "default_model": settings.default_model,
                "temperature": settings.temperature
            }
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )

@router.put("/workspaces/{workspace_id}/settings")
def update_workspace_settings(
    workspace_id: int,
    updates: WorkspaceSettingsUpdate = Body(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        workspace = db.query(Workspace).filter_by(id=workspace_id).first()
        if not workspace or workspace.owner_id != current_user.id:
            return format_response(
                status=False,
                message="Only owner can update settings",
                errors=[{"field": "permission", "message": "Only the workspace owner can update settings"}],
                status_code=403
            )

        settings = db.query(WorkspaceSettings).filter_by(workspace_id=workspace_id).first()
        if not settings:
            return format_response(
                status=False,
                message="Settings not found",
                errors=[{"field": "workspace_id", "message": "Workspace settings not found"}],
                status_code=404
            )

        if updates.default_voice is not None:
            settings.default_voice = updates.default_voice
        if updates.default_model is not None:
            settings.default_model = updates.default_model
        if updates.temperature is not None:
            settings.temperature = updates.temperature

        db.commit()

        return format_response(
            status=True,
            message="Settings updated"
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )
        
@router.get("/list-knowledge-bases")
def list_knowledge_bases(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # Join KnowledgeBase -> Workspace -> WorkspaceMember
        knowledge_bases = (
            db.query(KnowledgeBase)
            .join(Workspace, KnowledgeBase.workspace)  # join workspace from KB
            .outerjoin(WorkspaceMember, Workspace.id == WorkspaceMember.workspace_id)
            .filter(
                (Workspace.owner_id == current_user.id) |
                (WorkspaceMember.user_id == current_user.id)
            )
            .all()
        )

        result = []

        for kb in knowledge_bases:
            files = db.query(KnowledgeFile).filter_by(kb_id=kb.id).all()
            file_sources = []
            for file in files:
                file_sources.append({
                    "type": "document",
                    "source_id": str(file.id),
                    "filename": file.filename,
                    "file_url": file.file_url 
                })

            result.append({
                "knowledge_base_id": f"knowledge_base_{kb.id}",
                "knowledge_base_name": kb.name,
                "status": getattr(kb, "status", "in_progress"),
                "knowledge_base_sources": file_sources,
                "enable_auto_refresh": kb.enable_auto_refresh,
                "last_refreshed_timestamp": int(kb.last_refreshed.timestamp() * 1000) if kb.last_refreshed else 0
            })

        return format_response(
            status=True,
            message="Knowledge bases retrieved successfully",
            data=result
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}],
            status_code=500
        )
        
#  API KEYS      
@router.post("/workspaces/{workspace_id}/api-keys")
def create_api_key(
    workspace_id: int,
    payload: dict = Body(...),  # expecting: { "name": "my key name" }
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if not is_user_in_workspace(current_user.id, workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")

    workspace = db.query(Workspace).filter_by(id=workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="Workspace not found")

    key_value = generate_api_key()
    api_key = APIKey(
        workspace_id=workspace.id,
        name=payload["name"],
        key_value=key_value,
        created_by=current_user.id
    )
    db.add(api_key)
    db.commit()
    db.refresh(api_key)

    return format_response(
    status=True,
    message="API Key created successfully",
    data={
        "id": api_key.id,
        "name": api_key.name,
        "key_value": api_key.key_value,
        "last_edited_by": current_user.username,
        "updated_at": api_key.updated_at.isoformat()
    }
)

@router.get("/workspaces/{workspace_id}/api-keys")
def list_api_keys(
    workspace_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if not is_user_in_workspace(current_user.id, workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")

    keys = db.query(APIKey).filter_by(workspace_id=workspace_id).all()
    data = [{
        "id": key.id,
        "name": key.name,
        "key_value": key.key_value,
        "last_edited_by": current_user.username,
        "updated_at": key.updated_at.isoformat(),
        "is_webhook_key": key.is_webhook_key
    } for key in keys]

    return format_response(
        status=True,
        message="API keys fetched successfully",
        data=data
    )

@router.delete("/workspaces/{workspace_id}/api-keys/{key_id}")
def delete_api_key(
    workspace_id: int,
    key_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Check if the user is authorized in this workspace
    if not is_user_in_workspace(current_user.id, workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")

    # Fetch the specific API key under that workspace
    key = db.query(APIKey).filter_by(id=key_id, workspace_id=workspace_id).first()
    if not key:
        raise HTTPException(status_code=404, detail="API key not found in this workspace")

    if key.is_webhook_key:
        raise HTTPException(status_code=403, detail="Cannot delete webhook key. Create a new key and set as webhook key first before deleting this key.")

    db.delete(key)
    db.commit()

    return format_response(
        status=True,
        message="API key deleted successfully"
    )

@router.put("/api-keys/{key_id}/rename")
def update_api_key_name(
    key_id: int,
    name: str = Body(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    key = db.query(APIKey).filter_by(id=key_id).first()
    if not key:
        raise HTTPException(status_code=404, detail="API key not found")

    key.name = name
    key.updated_at = datetime.utcnow()
    key.created_by = current_user.id
    db.commit()
    return {"message": "API key renamed"}

@router.put("/api-keys/{key_id}/set-webhook")
def set_webhook_api_key(
    key_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    key = db.query(APIKey).filter_by(id=key_id).first()
    if not key:
        raise HTTPException(status_code=404, detail="API key not found")

    # Clear existing webhook key
    db.query(APIKey).filter(
        APIKey.workspace_id == key.workspace_id,
        APIKey.is_webhook_key == True
    ).update({"is_webhook_key": False})

    key.is_webhook_key = True
    db.commit()
    return {"message": "Webhook API key set"}

__all__ = ['router']

# Agent apis--------------------------------------------------

# @router.post("/agents")
# def create_agent(
#     payload: AgentCreate,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     if not is_user_in_workspace(current_user.id, payload.workspace_id, db):
#         raise HTTPException(status_code=403, detail="You do not have access to this workspace")

#     new_agent = pbx_ai_agent(
#         workspace_id=payload.workspace_id,
#         name=payload.agent_name,
#         voice_id=payload.voice_id,
#         voice_model=payload.voice_model,
#         fallback_voice_ids=payload.fallback_voice_ids,
#         voice_temperature=payload.voice_temperature,
#         voice_speed=payload.voice_speed,
#         volume=payload.volume,
#         responsiveness=payload.responsiveness,
#         interruption_sensitivity=payload.interruption_sensitivity,
#         enable_backchannel=payload.enable_backchannel,
#         backchannel_frequency=payload.backchannel_frequency,
#         backchannel_words=payload.backchannel_words,
#         reminder_trigger_ms=payload.reminder_trigger_ms,
#         reminder_max_count=payload.reminder_max_count,
#         ambient_sound=payload.ambient_sound,
#         ambient_sound_volume=payload.ambient_sound_volume,
#         language=payload.language,
#         webhook_url=payload.webhook_url,
#         boosted_keywords=payload.boosted_keywords,
#         opt_out_sensitive_data_storage=payload.opt_out_sensitive_data_storage,
#         opt_in_signed_url=payload.opt_in_signed_url,
#         pronunciation_dictionary=[
#             p.model_dump() for p in payload.pronunciation_dictionary
#         ] if payload.pronunciation_dictionary else None,
#         normalize_for_speech=payload.normalize_for_speech,
#         end_call_after_silence_ms=payload.end_call_after_silence_ms,
#         max_call_duration_ms=payload.max_call_duration_ms,
#         voicemail_option=payload.voicemail_option.model_dump() if payload.voicemail_option else None,
#         post_call_analysis_data=[
#             a.model_dump() for a in payload.post_call_analysis_data
#         ] if payload.post_call_analysis_data else None,
#         post_call_analysis_model=payload.post_call_analysis_model,
#         begin_message_delay_ms=payload.begin_message_delay_ms,
#         ring_duration_ms=payload.ring_duration_ms,
#         stt_mode=payload.stt_mode,
#         vocab_specialization=payload.vocab_specialization,
#         allow_user_dtmf=payload.allow_user_dtmf,
#         user_dtmf_options=payload.user_dtmf_options.model_dump() if payload.user_dtmf_options else None,
#         denoising_mode=payload.denoising_mode,
#         response_engine=payload.response_engine.model_dump() if payload.response_engine else None,
#         version=payload.version,
#         last_modification_timestamp=int(time.time() * 1000),
#     )



#     db.add(new_agent)
#     db.commit()
#     db.refresh(new_agent)

#     response_data = {
#         "agent_id": f"{new_agent.id}",
#         "last_modification_timestamp": new_agent.last_modification_timestamp,
#         "agent_name": new_agent.name,
#         "response_engine": new_agent.response_engine,
#         "language": new_agent.language,
#         "opt_out_sensitive_data_storage": new_agent.opt_out_sensitive_data_storage,
#         "opt_in_signed_url": new_agent.opt_in_signed_url,
#         "end_call_after_silence_ms": new_agent.end_call_after_silence_ms,
#         "version": new_agent.version,
#         "is_published": new_agent.is_published,
#         "post_call_analysis_model": new_agent.post_call_analysis_model,
#         "voice_id": new_agent.voice_id,
#         "voice_model": new_agent.voice_model,
#         "voice_temperature": new_agent.voice_temperature,
#         "voice_speed": new_agent.voice_speed,
#         "volume": new_agent.volume,
#         "enable_backchannel": new_agent.enable_backchannel,
#         "backchannel_frequency": new_agent.backchannel_frequency,
#         "reminder_trigger_ms": new_agent.reminder_trigger_ms,
#         "reminder_max_count": new_agent.reminder_max_count,
#         "max_call_duration_ms": new_agent.max_call_duration_ms,
#         "interruption_sensitivity": new_agent.interruption_sensitivity,
#         "ambient_sound_volume": new_agent.ambient_sound_volume,
#         "responsiveness": new_agent.responsiveness,
#         "normalize_for_speech": new_agent.normalize_for_speech,
#         "begin_message_delay_ms": new_agent.begin_message_delay_ms,
#         "ring_duration_ms": new_agent.ring_duration_ms,
#         "stt_mode": new_agent.stt_mode,
#         "allow_user_dtmf": new_agent.allow_user_dtmf,
#         "user_dtmf_options": new_agent.user_dtmf_options,
#         "denoising_mode": new_agent.denoising_mode
#     }

#     return {
#         "status": True,
#         "data": response_data
#     }

@router.post("/agents")
def create_agent(
    payload: AgentCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if not is_user_in_workspace(current_user.id, payload.workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")

    new_agent = pbx_ai_agent(
        workspace_id=payload.workspace_id,
        name=payload.agent_name,
        voice_id=payload.voice_id,
        voice_model=payload.voice_model,
        fallback_voice_ids=payload.fallback_voice_ids,
        voice_temperature=payload.voice_temperature,
        voice_speed=payload.voice_speed,
        volume=payload.volume,
        responsiveness=payload.responsiveness,
        interruption_sensitivity=payload.interruption_sensitivity,
        enable_backchannel=payload.enable_backchannel,
        backchannel_frequency=payload.backchannel_frequency,
        backchannel_words=payload.backchannel_words,
        reminder_trigger_ms=payload.reminder_trigger_ms,
        reminder_max_count=payload.reminder_max_count,
        ambient_sound=payload.ambient_sound,
        ambient_sound_volume=payload.ambient_sound_volume,
        language=payload.language,
        webhook_url=payload.webhook_url,
        boosted_keywords=payload.boosted_keywords,
        opt_out_sensitive_data_storage=payload.opt_out_sensitive_data_storage,
        opt_in_signed_url=payload.opt_in_signed_url,
        pronunciation_dictionary=[
            p.model_dump() for p in payload.pronunciation_dictionary
        ] if payload.pronunciation_dictionary else None,
        normalize_for_speech=payload.normalize_for_speech,
        end_call_after_silence_ms=payload.end_call_after_silence_ms,
        max_call_duration_ms=payload.max_call_duration_ms,
        post_call_analysis_data=[
            a.model_dump() for a in payload.post_call_analysis_data
        ] if payload.post_call_analysis_data else None,
        post_call_analysis_model=payload.post_call_analysis_model,
        begin_message_delay_ms=payload.begin_message_delay_ms,
        ring_duration_ms=payload.ring_duration_ms,
        stt_mode=payload.stt_mode,
        vocab_specialization=payload.vocab_specialization,
        allow_user_dtmf=payload.allow_user_dtmf,
        user_dtmf_options=payload.user_dtmf_options.model_dump() if payload.user_dtmf_options else None,
        denoising_mode=payload.denoising_mode,
        response_engine=payload.response_engine.model_dump() if payload.response_engine else None,
        version=payload.version,
        last_modification_timestamp=int(time.time() * 1000),
        voicemail_option=payload.voicemail_option.model_dump() if payload.voicemail_option else None,
        
    )

    db.add(new_agent)
    db.commit()
    db.refresh(new_agent)

    response_data = {
        "agent_id": str(new_agent.id),
        "last_modification_timestamp": new_agent.last_modification_timestamp,
        "agent_name": new_agent.name,
        "response_engine": new_agent.response_engine,
        "language": new_agent.language,
        "opt_out_sensitive_data_storage": new_agent.opt_out_sensitive_data_storage,
        "opt_in_signed_url": new_agent.opt_in_signed_url,
        "end_call_after_silence_ms": new_agent.end_call_after_silence_ms,
        "version": new_agent.version,
        "is_published": new_agent.is_published,
        "post_call_analysis_model": new_agent.post_call_analysis_model,
        "voice_id": new_agent.voice_id,
        "voice_model": new_agent.voice_model,
        "fallback_voice_ids": new_agent.fallback_voice_ids,
        "voice_temperature": new_agent.voice_temperature,
        "voice_speed": new_agent.voice_speed,
        "volume": new_agent.volume,
        "enable_backchannel": new_agent.enable_backchannel,
        "backchannel_frequency": new_agent.backchannel_frequency,
        "backchannel_words": new_agent.backchannel_words,
        "reminder_trigger_ms": new_agent.reminder_trigger_ms,
        "reminder_max_count": new_agent.reminder_max_count,
        "max_call_duration_ms": new_agent.max_call_duration_ms,
        "interruption_sensitivity": new_agent.interruption_sensitivity,
        "ambient_sound": new_agent.ambient_sound,
        "ambient_sound_volume": new_agent.ambient_sound_volume,
        "responsiveness": new_agent.responsiveness,
        "normalize_for_speech": new_agent.normalize_for_speech,
        "begin_message_delay_ms": new_agent.begin_message_delay_ms,
        "ring_duration_ms": new_agent.ring_duration_ms,
        "stt_mode": new_agent.stt_mode,
        "vocab_specialization": new_agent.vocab_specialization,
        "allow_user_dtmf": new_agent.allow_user_dtmf,
        "user_dtmf_options": new_agent.user_dtmf_options,
        "denoising_mode": new_agent.denoising_mode,
        "webhook_url": new_agent.webhook_url,
        "boosted_keywords": new_agent.boosted_keywords,
        "pronunciation_dictionary": new_agent.pronunciation_dictionary,
        "voicemail_option": new_agent.voicemail_option,
        "post_call_analysis_data": new_agent.post_call_analysis_data
    }

    return {
        "status": True,
        "data": response_data
    }


@router.delete("/agent/delete-agent/{agent_id}")
def delete_agent(
    agent_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    agent = db.query(pbx_ai_agent).filter(pbx_ai_agent.id == agent_id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")

    if not is_user_in_workspace(current_user.id, agent.workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")

    # 🗑️ Delete the agent
    db.delete(agent)
    db.commit()

    return {
        "status": True,
        "message": f"Agent deleted successfully",
        "data": None
    }

@router.get("/all-agents/{workspace_id}")
def list_my_agents(
    workspace_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if not is_user_in_workspace(current_user.id, workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")
    try:
        # Get all workspace IDs the user is a member of
        workspace_ids = db.query(WorkspaceMember.workspace_id).filter(
            WorkspaceMember.user_id == current_user.id
        ).subquery()

        # Fetch all agents in those workspaces
        agents = db.query(pbx_ai_agent).filter(
            pbx_ai_agent.workspace_id.in_(workspace_ids)
        ).all()

        # Serialize agents into your desired format
        agent_list = []
        for agent in agents:
            agent_list.append({
                "agent_id": agent.id,
                "version": agent.version,
                "is_published": agent.is_published,
                "response_engine": agent.response_engine or {},
                "agent_name": agent.name,
                "voice_id": agent.voice_id,
                "voice_model": agent.voice_model,
                "fallback_voice_ids": agent.fallback_voice_ids or [],
                "voice_temperature": agent.voice_temperature,
                "voice_speed": agent.voice_speed,
                "volume": agent.volume,
                "responsiveness": agent.responsiveness,
                "interruption_sensitivity": agent.interruption_sensitivity,
                "enable_backchannel": agent.enable_backchannel,
                "backchannel_frequency": agent.backchannel_frequency,
                "backchannel_words": agent.backchannel_words or [],
                "reminder_trigger_ms": agent.reminder_trigger_ms,
                "reminder_max_count": agent.reminder_max_count,
                "ambient_sound": agent.ambient_sound,
                "ambient_sound_volume": agent.ambient_sound_volume,
                "language": agent.language,
                "webhook_url": agent.webhook_url,
                "boosted_keywords": agent.boosted_keywords or [],
                "opt_out_sensitive_data_storage": agent.opt_out_sensitive_data_storage,
                "opt_in_signed_url": agent.opt_in_signed_url,
                "pronunciation_dictionary": agent.pronunciation_dictionary or [],
                "normalize_for_speech": agent.normalize_for_speech,
                "end_call_after_silence_ms": agent.end_call_after_silence_ms,
                "max_call_duration_ms": agent.max_call_duration_ms,
                "voicemail_option": agent.voicemail_option or {},
                "post_call_analysis_data": agent.post_call_analysis_data or [],
                "post_call_analysis_model": agent.post_call_analysis_model,
                "begin_message_delay_ms": agent.begin_message_delay_ms,
                "ring_duration_ms": agent.ring_duration_ms,
                "stt_mode": agent.stt_mode,
                "vocab_specialization": agent.vocab_specialization,
                "allow_user_dtmf": agent.allow_user_dtmf,
                "user_dtmf_options": agent.user_dtmf_options or {},
                "denoising_mode": agent.denoising_mode,
                "last_modification_timestamp": agent.last_modification_timestamp
            })

        return format_response(
            status=True,
            message="Agents fetched successfully",
            data=agent_list
        )

    except Exception as e:
        db.rollback()
        return format_response(
            status=False,
            message=f"Internal Server Error: {str(e)}",
            data=None
        )

@router.get("/agents/{agent_id}", response_model=APIResponse)
def get_agent(
    agent_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    agent = db.query(pbx_ai_agent).filter(pbx_ai_agent.id == agent_id).first()
    if not agent:
        return JSONResponse(
            status_code=404,
            content={
                "status": False,
                "message": "Agent not found",
                "data": None,
                "errors": [{"field": "agent_id", "message": "No agent found with this ID"}]
            }
        )

    # Optional: Check workspace access
    if not is_user_in_workspace(current_user.id, agent.workspace_id, db):
        return JSONResponse(
            status_code=403,
            content={
                "status": False,
                "message": "Access denied to this workspace",
                "data": None,
                "errors": [{"field": "workspace", "message": "You don't have access to this workspace"}]
            }
        )

    # Return agent data in desired format
    return JSONResponse(
        status_code=200,
        content={
            "status": True,
            "message": "Agent retrieved successfully",
            "data": AgentOut.from_orm(agent).model_dump()           
        }
    )

@router.patch("/agent/update-agent/{agent_id}")
def update_agent(
    agent_id: str,
    payload: AgentUpdate,  # ✅ using optional fields schema
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    agent = db.query(pbx_ai_agent).filter(pbx_ai_agent.id == agent_id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")
    if not is_user_in_workspace(current_user.id, agent.workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")

    update_data = payload.dict(exclude_unset=True)
    # Map `agent_name` in payload to `name` in DB model
    if "agent_name" in update_data:
        update_data["name"] = update_data.pop("agent_name")
    if not update_data:
        raise HTTPException(status_code=400, detail="No fields provided for update")

    for field, value in update_data.items():
        if hasattr(agent, field):
            setattr(agent, field, value)
        else:
            print(f"⚠️ Skipped unknown field: {field}")

    agent.last_modification_timestamp = int(time.time() * 1000)

    db.add(agent)  # ✅ Ensure it's tracked
    db.commit()
    db.refresh(agent)

    response_data = {
        "agent_id": str(agent.id),
        "last_modification_timestamp": agent.last_modification_timestamp,
        "agent_name": agent.name,
        "response_engine": agent.response_engine,
        "language": agent.language,
        "opt_out_sensitive_data_storage": agent.opt_out_sensitive_data_storage,
        "opt_in_signed_url": agent.opt_in_signed_url,
        "end_call_after_silence_ms": agent.end_call_after_silence_ms,
        "version": agent.version,
        "is_published": agent.is_published,
        "post_call_analysis_model": agent.post_call_analysis_model,
        "voice_id": agent.voice_id,
        "voice_model": agent.voice_model,
        "fallback_voice_ids": agent.fallback_voice_ids,
        "voice_temperature": agent.voice_temperature,
        "voice_speed": agent.voice_speed,
        "volume": agent.volume,
        "enable_backchannel": agent.enable_backchannel,
        "backchannel_frequency": agent.backchannel_frequency,
        "backchannel_words": agent.backchannel_words,
        "reminder_trigger_ms": agent.reminder_trigger_ms,
        "reminder_max_count": agent.reminder_max_count,
        "max_call_duration_ms": agent.max_call_duration_ms,
        "interruption_sensitivity": agent.interruption_sensitivity,
        "ambient_sound": agent.ambient_sound,
        "ambient_sound_volume": agent.ambient_sound_volume,
        "responsiveness": agent.responsiveness,
        "normalize_for_speech": agent.normalize_for_speech,
        "begin_message_delay_ms": agent.begin_message_delay_ms,
        "ring_duration_ms": agent.ring_duration_ms,
        "stt_mode": agent.stt_mode,
        "vocab_specialization": agent.vocab_specialization,
        "allow_user_dtmf": agent.allow_user_dtmf,
        "user_dtmf_options": agent.user_dtmf_options,
        "denoising_mode": agent.denoising_mode,
        "webhook_url": agent.webhook_url,
        "boosted_keywords": agent.boosted_keywords,
        "pronunciation_dictionary": agent.pronunciation_dictionary,
        "voicemail_option": agent.voicemail_option,
        "post_call_analysis_data": agent.post_call_analysis_data
    }

    return {
        "status": True,
        "message": "Agent updated successfully",
        "data": response_data
    }

# @router.patch("/agent/update-agent/{agent_id}")
# def update_agent(
#     agent_id: str,
#     payload: AgentCreate,  # Or create a separate AgentUpdate schema if fields are optional
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     # 🔥 Fetch agent
#     agent = db.query(pbx_ai_agent).filter(pbx_ai_agent.id == agent_id).first()
#     if not agent:
#         raise HTTPException(status_code=404, detail="Agent not found")

#     # ✅ Check workspace access
#     if not is_user_in_workspace(current_user.id, agent.workspace_id, db):
#         raise HTTPException(status_code=403, detail="You do not have access to this workspace")

#     # 🔄 Update only provided fields
#     update_data = payload.dict(exclude_unset=True)
#     for field, value in update_data.items():
#         setattr(agent, field, value)

#     agent.last_modification_timestamp = int(time.time() * 1000)

#     db.commit()
#     db.refresh(agent)

#     response_data = {
#         "agent_id": f"{agent.id}",
#         "last_modification_timestamp": agent.last_modification_timestamp,
#         "agent_name": agent.name,
#         "response_engine": agent.response_engine,
#         "language": agent.language,
#         "opt_out_sensitive_data_storage": agent.opt_out_sensitive_data_storage,
#         "opt_in_signed_url": agent.opt_in_signed_url,
#         "end_call_after_silence_ms": agent.end_call_after_silence_ms,
#         "version": agent.version,
#         "is_published": agent.is_published,
#         "post_call_analysis_model": agent.post_call_analysis_model,
#         "voice_id": agent.voice_id,
#         "voice_model": agent.voice_model,
#         "voice_temperature": agent.voice_temperature,
#         "voice_speed": agent.voice_speed,
#         "volume": agent.volume,
#         "enable_backchannel": agent.enable_backchannel,
#         "backchannel_frequency": agent.backchannel_frequency,
#         "reminder_trigger_ms": agent.reminder_trigger_ms,
#         "reminder_max_count": agent.reminder_max_count,
#         "max_call_duration_ms": agent.max_call_duration_ms,
#         "interruption_sensitivity": agent.interruption_sensitivity,
#         "ambient_sound_volume": agent.ambient_sound_volume,
#         "responsiveness": agent.responsiveness,
#         "normalize_for_speech": agent.normalize_for_speech,
#         "begin_message_delay_ms": agent.begin_message_delay_ms,
#         "ring_duration_ms": agent.ring_duration_ms,
#         "stt_mode": agent.stt_mode,
#         "allow_user_dtmf": agent.allow_user_dtmf,
#         "user_dtmf_options": agent.user_dtmf_options,
#         "denoising_mode": agent.denoising_mode
#     }

#     return {
#         "status": True,
#         "message": "Agent updated successfully",
#         "data": response_data
#     }
# PBX LLM APIs--------------------------------------------------

@router.post("/pbx-llms", response_model=PBXLLMOut)
def create_pbx_llm(
    payload: PBXLLMCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Validate workspace access
    if not is_user_in_workspace(current_user.id, payload.workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")

    workspace = db.query(Workspace).filter(Workspace.id == payload.workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="Workspace not found")

    llm = PBXLLM(
        workspace_id=payload.workspace_id,
        version=payload.version,
        model=payload.model,
        s2s_model=payload.s2s_model,
        model_temperature=payload.model_temperature,
        model_high_priority=payload.model_high_priority,
        tool_call_strict_mode=payload.tool_call_strict_mode,
        general_prompt=payload.general_prompt,
        general_tools=payload.general_tools,
        states=payload.states,
        starting_state=payload.starting_state,
        begin_message=payload.begin_message,
        default_dynamic_variables=payload.default_dynamic_variables,
        knowledge_base_ids=payload.knowledge_base_ids,
        last_modification_timestamp=int(time.time() * 1000)
    )
    db.add(llm)
    db.commit()
    db.refresh(llm)
    return PBXLLMOut(
        status=True,
        llm_id=f"{llm.id}"
    )
   
@router.delete("/delete-pbx-llm/{llm_id}")
def delete_pbx_llm(
    llm_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Fetch the PBXLLM
    llm = db.query(PBXLLM).filter(PBXLLM.id == llm_id).first()
    if not llm:
        return format_response(
            status=False,
            message="PBXLLM not found",
            errors=[{"field": "llm_id", "message": "PBXLLM not found"}],
            status_code=404
        )

    # Validate workspace access
    if not is_user_in_workspace(current_user.id, llm.workspace_id, db):
        return format_response(
            status=False,
            message="Access denied",
            errors=[{"field": "workspace", "message": "You do not have access to this workspace"}],
            status_code=403
        )

    # Delete PBXLLM
    db.delete(llm)
    db.commit()

    return format_response(
        status=True,
        message="LLM deleted successfully",
        data=None
    )


@router.get("/pbx-llm/{llm_id}")
def get_pbx_llm(llm_id: str, 
                current_user: User = Depends(get_current_user),
                db: Session = Depends(get_db)):
    # Fetch PBXLLM by id
    llm = db.query(PBXLLM).filter(PBXLLM.id == llm_id).first()

    if not llm:
        # Return 404 if not found
        return JSONResponse(
            status_code=404,
            content={
                "status": False,
                "message": f"LLM with id {llm_id} not found",
                "data": None
            }
        )

    # Build response data
    llm_data = {
        "llm_id": llm.id,
        "version": llm.version,
        "model": llm.model,
        "s2s_model": llm.s2s_model,
        "model_temperature": llm.model_temperature,
        "model_high_priority": llm.model_high_priority,
        "tool_call_strict_mode": llm.tool_call_strict_mode,
        "general_prompt": llm.general_prompt,
        "general_tools": llm.general_tools,
        "states": llm.states,
        "starting_state": llm.starting_state,
        "begin_message": llm.begin_message,
        "default_dynamic_variables": llm.default_dynamic_variables,
        "knowledge_base_ids": llm.knowledge_base_ids,
        "last_modification_timestamp": llm.last_modification_timestamp,
        "is_published": llm.is_published
    }

    # Return wrapped response
    return JSONResponse(
        status_code=200,
        content={
            "status": True,
            "message": "",
            "data": llm_data
        }
    )

@router.patch("/pbx-llms/update-llm/{llm_id}")
def update_pbx_llm(
    llm_id: str,
    payload: PBXLLMCreate,  # Or make a PBXLLMUpdate schema where all fields are Optional
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 🔥 Fetch PBXLLM by ID
    llm = db.query(PBXLLM).filter(PBXLLM.id == llm_id).first()
    if not llm:
        raise HTTPException(status_code=404, detail=f"LLM with id {llm_id} not found")

    # ✅ Check workspace access
    if not is_user_in_workspace(current_user.id, llm.workspace_id, db):
        raise HTTPException(status_code=403, detail="You do not have access to this workspace")

    # 🔄 Update only provided fields
    update_data = payload.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(llm, field, value)

    # Update timestamp
    llm.last_modification_timestamp = int(time.time() * 1000)

    db.commit()
    db.refresh(llm)

    # Build response
    updated_llm_data = {
        "llm_id": llm.id,
        "version": llm.version,
        "model": llm.model,
        "s2s_model": llm.s2s_model,
        "model_temperature": llm.model_temperature,
        "model_high_priority": llm.model_high_priority,
        "tool_call_strict_mode": llm.tool_call_strict_mode,
        "general_prompt": llm.general_prompt,
        "general_tools": llm.general_tools,
        "states": llm.states,
        "starting_state": llm.starting_state,
        "begin_message": llm.begin_message,
        "default_dynamic_variables": llm.default_dynamic_variables,
        "knowledge_base_ids": llm.knowledge_base_ids,
        "last_modification_timestamp": llm.last_modification_timestamp,
        "is_published": llm.is_published
    }

    return JSONResponse(
        status_code=200,
        content={
            "status": True,
            "message": "LLM updated successfully",
            "data": updated_llm_data
        }
    )

@router.get("/all-pbx-llms/{workspace_id}")
def get_pbx_llm(
    workspace_id=int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        # Fetch the PBX LLM for the user's workspace
        pbx_llm = (
            db.query(PBXLLM)
            .filter(PBXLLM.workspace_id == workspace_id)            
            .all()
        )

        if not pbx_llm:
            return format_response(
                status=False,
                message="No LLM configuration found for this workspace",
                data=None,
                errors=[{"field": "workspace_id", "message": "No PBX LLM found for workspace"}],
                status_code=404
            )
       
        # Validate and serialize with Pydantic
        pbx_llm_out_list = [GetPBXLLMOut.from_orm(llm).model_dump() for llm in pbx_llm] 

        return format_response(
            status=True,
            message="LLM configuration fetched successfully",
            data=pbx_llm_out_list
        )

    except Exception as e:
        return format_response(
            status=False,
            message="Failed to fetch LLM configuration",
            errors=[{"field": "server", "message": str(e)}],
            status_code=500
        )

# Chat room APIs--------------------------------------------------
@router.post("/create-chat", response_model=CreateChatResponse)
def create_chat(payload: CreateChatRequest, db: Session = Depends(get_db)):
    try:
        chat = ChatSession(
            agent_id=payload.agent_id,
            agent_version=payload.agent_version or 0,
            chat_status=payload.chat_status,
            retell_llm_dynamic_variables=payload.retell_llm_dynamic_variables,
            collected_dynamic_variables={},
            chat_metadata=payload.metadata,
            start_timestamp=payload.start_timestamp,
            end_timestamp=payload.end_timestamp,
            transcript=payload.transcript or "null",
            message_with_tool_calls=[],  # ensure it's a list of dicts or JSON-serializable
            chat_cost=payload.chat_cost.model_dump() if payload.chat_cost else None,
            chat_analysis=payload.chat_analysis.model_dump() if payload.chat_analysis else None
        )

        db.add(chat)
        db.commit()
        db.refresh(chat)

        return CreateChatResponse(
            chat_id=chat.chat_id,
            agent_id=chat.agent_id,
            chat_status=chat.chat_status,
            retell_llm_dynamic_variables=chat.retell_llm_dynamic_variables,
            collected_dynamic_variables=chat.collected_dynamic_variables,
            start_timestamp=chat.start_timestamp,
            end_timestamp=chat.end_timestamp,
            transcript=chat.transcript,
            message_with_tool_calls=chat.message_with_tool_calls,
            metadata=chat.chat_metadata,
            chat_cost=chat.chat_cost,
            chat_analysis=chat.chat_analysis
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
# voice
@router.post("/llm-voice", response_model=VoiceOut)
def create_voice(
    payload: VoiceCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)  # optional if you want auth
):
    existing = db.query(LLMVoice).filter_by(voice_id=payload.voice_id).first()
    if existing:
        raise HTTPException(status_code=400, detail="Voice ID already exists")
    
    voice = LLMVoice(**payload.dict())
    db.add(voice)
    db.commit()
    db.refresh(voice)
    return voice

@router.get("/llm-voice/{voice_id}", response_model=VoiceOut)
def get_voice(
    voice_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)  # optional
):
    voice = db.query(LLMVoice).filter_by(voice_id=voice_id).first()
    if not voice:
        raise HTTPException(status_code=404, detail="Voice not found")
    return voice

@router.get("/all-voices", response_model=APIResponse)
def list_voices(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        voices_data = [
            {
                "voice_id": "openai-shimmer",
                "voice_name": "Shimmer",
                "provider": "elevenlabs",
                "gender": "female",
                "accent": "Neutral",
                "age": "Adult",
                "preview_audio_url": "https://example.com/previews/shimmer.mp3"
            },
            {
                "voice_id": "elevn-echo",
                "voice_name": "Echo",
                "provider": "elevenlabs",
                "gender": "male",
                "accent": "British",
                "age": "Middle-aged",
                "preview_audio_url": "https://example.com/previews/echo.mp3"
            },
            {
                "voice_id": "elevn-nova",
                "voice_name": "Nova",
                "provider": "elevenlabs",
                "gender": "female",
                "accent": "American",
                "age": "Young Adult",
                "preview_audio_url": "https://example.com/previews/nova.mp3"
            },
            {
                "voice_id": "elevn-pulse",
                "voice_name": "Pulse",
                "provider": "elevenlabs",
                "gender": "male",
                "accent": "Australian",
                "age": "Adult",
                "preview_audio_url": "https://example.com/previews/pulse.mp3"
            }
        ]
        
        # Validate and convert to VoiceOut objects
        try:
            voices = [VoiceOut.model_validate(voice) for voice in voices_data]
        except ValueError as validation_error:
            return APIResponse(
                status=False,
                message="Validation error",
                data=None,
                errors=[
                    {
                        "field": "voice_data",
                        "message": str(validation_error)
                    }
                ]
            )
        
        # Return successful response
        return APIResponse(
            status=True,
            message="Voices fetched successfully",
            data=voices,
            errors=[]
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail={
                "status": False,
                "message": "Internal Server Error",
                "data": None,
                "errors": [
                    {
                        "field": "server",
                        "message": str(e)
                    }
                ]
            }
        )

#import phone
@router.post("/import-phone-number", response_model=PhoneNumberOut, status_code=201)
def import_phone_number(
    payload: PhoneNumberCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)  # optional
):
    existing = db.query(ImportedPhoneNumber).filter_by(phone_number=payload.phone_number).first()
    if existing:
        raise HTTPException(status_code=400, detail="Phone number already exists")

    area_code = int(payload.phone_number[2:5])
    pretty = f"+1 ({payload.phone_number[2:5]}) {payload.phone_number[5:8]}-{payload.phone_number[8:]}"
    timestamp = int(time.time() * 1000)

    phone = ImportedPhoneNumber(
        phone_number=payload.phone_number,
        phone_number_type=payload.phone_number_type,
        phone_number_pretty=pretty,
        inbound_agent_id=payload.inbound_agent_id,
        outbound_agent_id=payload.outbound_agent_id,
        inbound_agent_version=payload.inbound_agent_version,
        outbound_agent_version=payload.outbound_agent_version,
        area_code=area_code,
        nickname=payload.nickname,
        inbound_webhook_url=payload.inbound_webhook_url,
        last_modification_timestamp=timestamp,
    )

    db.add(phone)
    db.commit()
    db.refresh(phone)
    return phone

@router.get("/import-phone-number/{phone_number}", response_model=PhoneNumberOut)
def get_phone_number(
    phone_number: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)  # optional
):
    phone = db.query(ImportedPhoneNumber).filter_by(phone_number=phone_number).first()
    if not phone:
        raise HTTPException(status_code=404, detail="Phone number not found")
    return phone

@router.get("/all-import-phone-number", response_model=List[PhoneNumberOut])
def list_phone_numbers(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)  # optional
):
    return db.query(ImportedPhoneNumber).all()

@router.get("/list-phone-numbers")
def list_phone_numbers(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        phone_numbers_data = [
            {
                "id": "pn_001",
                "phone_number": "+14157774444",
                "phone_number_type": "retell-twilio",
                "phone_number_pretty": "+1 (415) 777-4444",
                "inbound_agent_id": "oBeDLoLOeuAbiuaMFXRtDOLriTJ5tSxD",
                "outbound_agent_id": "oBeDLoLOeuAbiuaMFXRtDOLriTJ5tSxD",
                "inbound_agent_version": "v1.0",
                "outbound_agent_version": "v1.0",
                "area_code": "415",
                "nickname": "Main Line",
                "inbound_webhook_url": "https://api.example.com/webhooks/inbound/main",
                "last_modification_timestamp": "2025-07-07T18:45:00+05:30"
            },
            {
                "id": "pn_002",
                "phone_number": "+12025550123",
                "phone_number_type": "retell-twilio",
                "phone_number_pretty": "+1 (202) 555-0123",
                "inbound_agent_id": "xYcZMnPQrsTuvWxyZABcDEFghIJkLmNo",
                "outbound_agent_id": "xYcZMnPQrsTuvWxyZABcDEFghIJkLmNo",
                "inbound_agent_version": "v1.1",
                "outbound_agent_version": "v1.1",
                "area_code": "202",
                "nickname": "Support Line",
                "inbound_webhook_url": "https://api.example.com/webhooks/inbound/support",
                "last_modification_timestamp": "2025-07-06T12:30:00+05:30"
            },
            {
                "id": "pn_003",
                "phone_number": "+13105556789",
                "phone_number_type": "retell-twilio",
                "phone_number_pretty": "+1 (310) 555-6789",
                "inbound_agent_id": "pQrStUvWxYzAbCdEfGhIjKlMnOpQrStU",
                "outbound_agent_id": "pQrStUvWxYzAbCdEfGhIjKlMnOpQrStU",
                "inbound_agent_version": "v2.0",
                "outbound_agent_version": "v2.0",
                "area_code": "310",
                "nickname": "Sales Line",
                "inbound_webhook_url": "https://api.example.com/webhooks/inbound/sales",
                "last_modification_timestamp": "2025-07-05T09:15:00+05:30"
            },
            {
                "id": "pn_004",
                "phone_number": "+14405559876",
                "phone_number_type": "retell-twilio",
                "phone_number_pretty": "+1 (440) 555-9876",
                "inbound_agent_id": "aBcDeFgHiJkLmNoPqRsTuVwXyZaBcDeF",
                "outbound_agent_id": "aBcDeFgHiJkLmNoPqRsTuVwXyZaBcDeF",
                "inbound_agent_version": "v1.2",
                "outbound_agent_version": "v1.2",
                "area_code": "440",
                "nickname": "Customer Care",
                "inbound_webhook_url": "https://api.example.com/webhooks/inbound/customer",
                "last_modification_timestamp": "2025-07-04T14:20:00+05:30"
            }
        ]

        # Validate data with PhoneNumberOut schema
        data = [PhoneNumberOut(**phone).dict() for phone in phone_numbers_data]

        return format_response(
            status=True,
            message="Phone numbers retrieved successfully",
            data=data
        )

    except Exception as e:
        return format_response(
            status=False,
            message="Internal Server Error",
            errors=[{"field": "server", "message": str(e)}]
        )
    
# -----------------web call api--------------------------------------------------

@router.post("/call/create-web-call", response_model=WebCallResponse)
async def create_web_call(
    payload: WebCallCreateRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Fetch agent details from pbx_ai_agent table
    agent = db.query(pbx_ai_agent).filter(pbx_ai_agent.id == payload.agent_id).first()
   
    if not agent:
        return JSONResponse(
            status_code=404,
            content={
                "status": False,
                "message": "Agent not found",
                "data": None
            }
        )
    # Generate unique call_id and access_token
    call_id = uuid.uuid4().hex
    print(f"Creating web call with ID 🔥: {call_id}")
    # 🔥 LiveKit credentials from .env
    LIVEKIT_API_KEY = os.getenv("LIVEKIT_API_KEY")
    LIVEKIT_API_SECRET = os.getenv("LIVEKIT_API_SECRET")
    LIVEKIT_URL = os.getenv("LIVEKIT_URL")

    if not (LIVEKIT_API_KEY and LIVEKIT_API_SECRET and LIVEKIT_URL):
        raise HTTPException(status_code=500, detail="LiveKit credentials not configured")

    # 🔥 Create LiveKit Room
    lkapi = api.LiveKitAPI(
        url=LIVEKIT_URL,
        api_key=LIVEKIT_API_KEY,
        api_secret=LIVEKIT_API_SECRET
    )
    try:
        lkapi.room.create_room(
            api.CreateRoomRequest(
                name=call_id,  # Use call_id as room name
                empty_timeout=10 * 60,  # 10 minutes timeout
                max_participants=10
            )
        )
        print(f"LiveKit Room '{call_id}' created successfully.")
    except Exception as e:
        if "already exists" not in str(e):
            raise HTTPException(status_code=500, detail=f"LiveKit Error: {e}")

    # 🔥 Generate LiveKit Access Token
    token = api.AccessToken(LIVEKIT_API_KEY, LIVEKIT_API_SECRET) \
        .with_identity(f"user_{current_user.id}") \
        .with_name(current_user.username or f"User {current_user.id}") \
        .with_grants(api.VideoGrants(
            room_join=True,
            room=call_id
        ))
   
    # Create WebCall DB record
    new_web_call = WebCall(
        call_type="web_call",
        call_id=call_id,  # ✅ Set call_id here
        access_token=token.to_jwt(),
        agent_id=agent.id,
        agent_version=agent.version,
        call_status="registered",
        call_metadata={
        "livekit_url": LIVEKIT_URL
        },
        retell_llm_dynamic_variables={},
        collected_dynamic_variables={},
        custom_sip_headers={},
        latency={},
        call_cost={
            "product_costs": [],
            "total_duration_unit_price": 0,
            "total_duration_seconds": 0,
            "total_one_time_price": 0,
            "combined_cost": 0
        },
        opt_out_sensitive_data_storage=agent.opt_out_sensitive_data_storage,
        opt_in_signed_url=agent.opt_in_signed_url,
        start_timestamp=None,
        end_timestamp=None,
        duration_ms=None,
        transcript=None,
        transcript_object={},
        transcript_with_tool_calls=[],
        recording_url=None,
        public_log_url=None,
        knowledge_base_retrieved_contents_url=None,
        disconnection_reason=None,
        call_analysis=None,
        llm_token_usage=None,
        created_at=int(time.time() * 1000),
        updated_at=int(time.time() * 1000)
    )

    db.add(new_web_call)
    db.commit()
    db.refresh(new_web_call)

    # 🔥 Dispatch agent to the room
    try:
        dispatch = await lkapi.agent_dispatch.create_dispatch(
            api.CreateAgentDispatchRequest(
                room=call_id,
                agent_name="ankit",
                metadata=json.dumps({
                    "agent_id": agent.id
                })
            )
        )
        print(f"✅ Agent dispatched to room '{call_id}' with Dispatch ID: {dispatch.id}")
    except Exception as e:
        print(f"⚠️ Failed to dispatch agent: {e}")
        raise HTTPException(status_code=500, detail=f"Dispatch Error: {e}")
    finally:
        await lkapi.aclose()

    data = WebCallResponse(
        call_type=new_web_call.call_type,
        access_token=new_web_call.access_token,
        call_id=new_web_call.call_id,
        agent_id=new_web_call.agent_id,
        agent_name=agent.name,
        agent_version=new_web_call.agent_version,
        call_status=new_web_call.call_status,
        call_metadata=new_web_call.call_metadata,
        retell_llm_dynamic_variables=new_web_call.retell_llm_dynamic_variables,
        collected_dynamic_variables=new_web_call.collected_dynamic_variables,
        custom_sip_headers=new_web_call.custom_sip_headers,
        opt_out_sensitive_data_storage=new_web_call.opt_out_sensitive_data_storage,
        opt_in_signed_url=new_web_call.opt_in_signed_url,
        start_timestamp=new_web_call.start_timestamp,
        end_timestamp=new_web_call.end_timestamp,
        duration_ms=new_web_call.duration_ms,
        transcript=new_web_call.transcript,
        transcript_object=new_web_call.transcript_object,
        transcript_with_tool_calls=new_web_call.transcript_with_tool_calls,
        recording_url=new_web_call.recording_url,
        public_log_url=new_web_call.public_log_url,
        knowledge_base_retrieved_contents_url=new_web_call.knowledge_base_retrieved_contents_url,
        latency=new_web_call.latency,
        disconnection_reason=new_web_call.disconnection_reason,
        call_analysis=new_web_call.call_analysis,
        call_cost=new_web_call.call_cost,
        llm_token_usage=new_web_call.llm_token_usage
    ).dict()

    return JSONResponse(
        status_code=200,
        content={
            "status": True,
            "message": "Web call created successfully",
            "data": data
        }
    )

@router.get("/get-webcall/{call_id}", response_model=WebCallResponse)
def get_web_call_by_id(
    call_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Fetch WebCall record from DB
    web_call = db.query(WebCall).filter(WebCall.call_id == call_id).first()
    # Fetch agent details from pbx_ai_agent table
    agent = db.query(pbx_ai_agent).filter(pbx_ai_agent.id == web_call.agent_id).first()
   
    if not agent:
        return JSONResponse(
            status_code=404,
            content={
                "status": False,
                "message": "Agent not found",
                "data": None
            }
        )
    
    if not web_call:
        return JSONResponse(
            status_code=404,
            content={
                "status": False,
                "message": "WebCall not found",
                "data": None
            }
        )
       

    # Prepare response
    data = WebCallResponse(
        call_type=web_call.call_type,
        access_token=web_call.access_token,
        call_id=web_call.call_id,
        agent_id=web_call.agent_id,
        agent_name=agent.name if agent.name else None,  # Assuming relationship
        agent_version=web_call.agent_version,
        call_status=web_call.call_status,
        call_metadata=web_call.call_metadata,
        retell_llm_dynamic_variables=web_call.retell_llm_dynamic_variables,
        collected_dynamic_variables=web_call.collected_dynamic_variables,
        custom_sip_headers=web_call.custom_sip_headers,
        opt_out_sensitive_data_storage=web_call.opt_out_sensitive_data_storage,
        opt_in_signed_url=web_call.opt_in_signed_url,
        start_timestamp=web_call.start_timestamp,
        end_timestamp=web_call.end_timestamp,
        duration_ms=web_call.duration_ms,
        transcript=web_call.transcript,
        transcript_object=web_call.transcript_object,
        transcript_with_tool_calls=web_call.transcript_with_tool_calls,
        recording_url=web_call.recording_url,
        public_log_url=web_call.public_log_url,
        knowledge_base_retrieved_contents_url=web_call.knowledge_base_retrieved_contents_url,
        latency=web_call.latency,
        disconnection_reason=web_call.disconnection_reason,
        call_analysis=web_call.call_analysis,
        call_cost=web_call.call_cost,
        llm_token_usage=web_call.llm_token_usage
    ).dict()

    return JSONResponse(
        status_code=200,
        content={
            "status": True,
            "message": "WebCall fetched successfully",
            "data": data
        }
    )

# import os
# from twilio.rest import Client

# account_sid = os.environ["AC0d8d96e8bec0d573804a160bd96483b3"]
# auth_token = os.environ["d55c8e547b69bbcb34cfc27bc7bc155f"]
# client = Client(account_sid, auth_token)

# incoming_phone_number = client.incoming_phone_numbers.create(
#     phone_number="+14155552344"
# )

# print(incoming_phone_number.account_sid)

from twilio.base.exceptions import TwilioRestException

# Twilio credentials from environment variables
TWILIO_ACCOUNT_SID = os.getenv("TWILIO_ACCOUNT_SID")
TWILIO_AUTH_TOKEN = os.getenv("TWILIO_AUTH_TOKEN")

client = Client(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN)

@router.post("/purchase-phone-number", response_model=Response)
def purchase_phone_number(
    data: PhoneNumberRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # Step 1: Search for available phone numbers
        available_numbers = client.available_phone_numbers("US").local.list(
            sms_enabled=True,
            voice_enabled=True,
            area_code=data.area_code
        )

        if not available_numbers:
            return JSONResponse(
                status_code=400,
                content={
                    "status": False,
                    "message": f"No available phone numbers found in area code {data.area_code}",
                    "data": None,
                    "errors": [{"field": "phone_number", "message": "No available numbers"}]
                }
            )

        # Step 2: Select and purchase number
        phone_number = available_numbers[0].phone_number
        purchased_number = client.incoming_phone_numbers.create(phone_number=phone_number)

        # Step 3: Format number
        national_number = phone_number[2:]  # remove '+1'
        area_code = int(national_number[:3])
        phone_number_pretty = f"+1 ({national_number[:3]}) {national_number[3:6]}-{national_number[6:]}"

        # Step 4: Save to DB
        phone_number_db = PhoneNumber(
            phone_number=purchased_number.phone_number,
            phone_number_type="ucaas-twilio",
            phone_number_pretty=phone_number_pretty,
            area_code=area_code,
            inbound_agent_id=data.inbound_agent_id,
            outbound_agent_id=data.outbound_agent_id,
            inbound_agent_version=data.inbound_agent_version or 1,
            outbound_agent_version=1,
            nickname=data.nickname,
            inbound_webhook_url=data.inbound_webhook_url or "https://example.com/inbound-webhook",
            last_modification_timestamp=1703413636133,
            owner_id=current_user.id
        )
        db.add(phone_number_db)

        if data.inbound_webhook_url:
            client.incoming_phone_numbers(purchased_number.sid).update(
                sms_url=data.inbound_webhook_url
            )

        db.commit()
        db.refresh(phone_number_db)

        # Step 5: Return success
        return {
            "status": True,
            "message": "Phone number purchased",
            "data": {
                "phone_number": phone_number_db.phone_number,
                "phone_number_type": phone_number_db.phone_number_type,
                "phone_number_pretty": phone_number_db.phone_number_pretty,
                "area_code": phone_number_db.area_code,
                "inbound_agent_id": data.inbound_agent_id,
                "inbound_agent_name": data.inbound_agent_name,
                "inbound_agent_version": data.inbound_agent_version,
                "inbound_webhook_url": data.inbound_webhook_url,
                "nickname": data.nickname,
                "number_provider": data.number_provider,
                "outbound_agent_id": data.outbound_agent_id,
                "outbound_agent_version": 1,
                "last_modification_timestamp": phone_number_db.last_modification_timestamp
            },
            "errors": None
        }

    except TwilioRestException as e:
        db.rollback()
        return JSONResponse(
            status_code=502,
            content={
                "status": False,
                "message": "Twilio API error",
                "data": None,
                "errors": [{"field": "twilio", "message": str(e)}]
            }
        )

    except Exception as e:
        db.rollback()
        return {
            "status": False,
            "message": "Internal Server Error",
            "data": None,
            "errors": [{"field": "server", "message": str(e)}]
        }
        
# delete the purchased phone number from twilio
@router.delete("/delete-phone-number/{phone_number}")
def delete_phone_number(
    phone_number: str,
    authorization: str = Header(..., description="Bearer token"),
    db: Session = Depends(get_db)
):
    try:
        # Step 1: Authenticate (this is a placeholder, implement real API key check here)
        if not authorization.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Invalid authorization header format")
        
        token = authorization.split(" ")[1]
        if token != "YOUR_API_KEY":  # Replace with real token check logic
            raise HTTPException(status_code=403, detail="Invalid or expired API key")

        # Step 2: Look up the phone number SID via Twilio
        incoming_numbers = client.incoming_phone_numbers.list(phone_number=phone_number)

        if not incoming_numbers:
            return {
                "status": False,
                "message": "Phone number not found in Twilio",
                "data": None,
                "errors": [{"field": "phone_number", "message": "Not found in Twilio"}]
            }

        number_sid = incoming_numbers[0].sid

        # Step 3: Delete from Twilio
        client.incoming_phone_numbers(number_sid).delete()

        # Step 4: Delete from local DB
        db_number = db.query(PhoneNumber).filter_by(phone_number=phone_number).first()
        if db_number:
            db.delete(db_number)
            db.commit()

        return {
            "status": True,
            "message": "Phone number deleted successfully",
            "data": {
                "phone_number": phone_number,
                "deleted_from_twilio": True,
                "deleted_from_database": db_number is not None
            },
            "errors": None
        }

    except TwilioRestException as e:
        db.rollback()
        return {
            "status": False,
            "message": "Twilio API error",
            "data": None,
            "errors": [{"field": "twilio", "message": str(e)}]
        }

    except Exception as e:
        db.rollback()
        return {
            "status": False,
            "message": "Internal Server Error",
            "data": None,
            "errors": [{"field": "server", "message": str(e)}]
        }


import logging

@router.post("/twilio/voice")
async def handle_incoming_call():
    """
    Twilio hits this endpoint when a call comes in.
    This responds with TwiML to forward the call to your SIP trunk (LiveKit).
    """
    sip_uri = "sip:2tfpk869du5.sip.livekit.cloud"
    twiml = f"""
    <Response>
        <Dial>
            <Sip>{sip_uri}</Sip>
        </Dial>
    </Response>
    """
    logging.info("Returning SIP TwiML to Twilio.")
    return Response(content=twiml, media_type="application/xml")

# optional 
@router.post("/livekit/webhook")
async def livekit_webhook(request: Request):
    body = await request.json()
    event = body.get("event")
    participant = body.get("participant", {}).get("identity")
    room = body.get("room", {}).get("name")

    logging.info(f"LiveKit event received: {event}")
    logging.info(f"Participant: {participant}, Room: {room}")

    if event == "participant_connected":
        logging.info(f"SIP participant {participant} joined room {room}. Notify agent or connect.")

    return {"status": "ok"}